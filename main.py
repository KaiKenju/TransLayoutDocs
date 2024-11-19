
# import os
# import shutil
# from pdf2image import convert_from_path
# from paddleocr import PPStructure, save_structure_res
# from ppstructure.recovery.recovery_to_doc import sorted_layout_boxes, convert_info_docx
# from docxcompose.composer import Composer
# from docx import Document
# import cv2
# import numpy as np
# from tqdm import tqdm

# # Khởi tạo PaddleOCR với ngôn ngữ tiếng Anh
# table_engine = PPStructure(recovery=True, lang='en')

# # Đường dẫn thư mục lưu trữ
# save_folder = './store_data'
# output_file = './outputs/'
# os.makedirs(save_folder, exist_ok=True)

# # Hàm xử lý từng trang PDF
# def process_image_to_word(img, page_number):
#     result = table_engine(img)
#     save_structure_res(result, save_folder, f'page_{page_number}')
    
#     for line in result:
#         line.pop('img')
#         print(line)
    
#     h, w, _ = img.shape
#     res = sorted_layout_boxes(result, w)
#     convert_info_docx(img, res, save_folder, f'page_{page_number}')
#     print(f"Trang {page_number} đã xử lý xong.")

# # Hàm gộp các file Word và đảm bảo ảnh hiển thị đúng
# def merge_word_documents_with_images(doc_folder, output_file):
#     # Lấy danh sách file .docx và sắp xếp theo thứ tự
#     doc_files = sorted(
#         [f for f in os.listdir(doc_folder) if f.endswith('.docx')],
#         key=lambda x: int(x.split('_')[1].split('.')[0]) if x.startswith('page_') else float('inf')
#     )
    
#     if not doc_files:
#         raise ValueError(f"Thư mục '{doc_folder}' không chứa file .docx hợp lệ.")
    
#     # Tạo file Word hợp nhất
#     master_doc = Document(os.path.join(doc_folder, doc_files[0]))
#     composer = Composer(master_doc)
    
#     # Hợp nhất các file còn lại
#     for filename in doc_files[1:]:
#         sub_doc_path = os.path.join(doc_folder, filename)
#         sub_doc = Document(sub_doc_path)
#         composer.append(sub_doc)
    
#     # Di chuyển tất cả thư mục "media" từ các file con vào gốc
#     for folder in os.listdir(doc_folder):
#         folder_path = os.path.join(doc_folder, folder)
#         if os.path.isdir(folder_path) and folder.startswith('page_'):
#             media_folder = os.path.join(folder_path, 'media')
#             if os.path.exists(media_folder):
#                 for media_file in os.listdir(media_folder):
#                     shutil.move(os.path.join(media_folder, media_file), os.path.join(doc_folder, media_file))
#                 shutil.rmtree(media_folder)
    
#     # Lưu file hợp nhất
#     composer.save(output_file)
#     print(f"Hợp nhất hoàn tất: {output_file}")

# # Xử lý PDF thành file Word
# def process_pdf_to_word(pdf_path):
#     # Chuyển đổi từng trang PDF thành hình ảnh
#     images = convert_from_path(pdf_path)
    
#     for i, image in enumerate(tqdm(images, desc="Pages processing")):
#         # Chuyển đổi từ PIL Image sang OpenCV Image
#         img = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
#         process_image_to_word(img, i + 1)
    
#     # Gộp các file Word
#     merge_word_documents_with_images(save_folder, './outputs/final_output.docx')
#     print("Hoàn thành! File Word hợp nhất lưu tại: ./outputs/final_output.docx")

# # Đường dẫn file PDF đầu vào
# pdf_path = './inputs/example1.pdf'
# process_pdf_to_word(pdf_path)


import os
import shutil
from pdf2image import convert_from_path
from paddleocr import PPStructure, save_structure_res
from Recovery.recovery_to_doc import sorted_layout_boxes, convert_info_docx
from docxcompose.composer import Composer
from docx import Document
import cv2
import numpy as np
from tqdm import tqdm

# Khởi tạo PaddleOCR với ngôn ngữ tiếng Anh
table_engine = PPStructure(recovery=True, lang='en')

# Đường dẫn thư mục lưu trữ
save_folder = './temp'
output_folder = './outputs/'
os.makedirs(save_folder, exist_ok=True)
os.makedirs(output_folder, exist_ok=True)

# Hàm xóa dữ liệu trong thư mục lưu trữ
def clear_folder(folder_path):
    if os.path.exists(folder_path):
        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)
            try:
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)
            except Exception as e:
                print(f"Lỗi khi xóa file hoặc thư mục {file_path}: {e}")

# Hàm xử lý từng trang PDF
def process_image_to_word(img, page_number):
    result = table_engine(img)
    save_structure_res(result, save_folder, f'page_{page_number}')
    
    for line in result:
        line.pop('img')
        print(line)
    
    h, w, _ = img.shape
    res = sorted_layout_boxes(result, w)
    convert_info_docx(img, res, save_folder, f'page_{page_number}')
    print(f"Trang {page_number} đã xử lý xong.")

# Hàm gộp các file Word và đảm bảo ảnh hiển thị đúng
def merge_word_documents_with_images(doc_folder, output_file):
    # Lấy danh sách file .docx và sắp xếp theo thứ tự
    doc_files = sorted(
        [f for f in os.listdir(doc_folder) if f.endswith('.docx')],
        key=lambda x: int(x.split('_')[1].split('.')[0]) if x.startswith('page_') else float('inf')
    )
    
    if not doc_files:
        raise ValueError(f"Thư mục '{doc_folder}' không chứa file .docx hợp lệ.")
    
    # Tạo file Word hợp nhất
    master_doc = Document(os.path.join(doc_folder, doc_files[0]))
    composer = Composer(master_doc)
    
    # Hợp nhất các file còn lại
    for filename in doc_files[1:]:
        sub_doc_path = os.path.join(doc_folder, filename)
        sub_doc = Document(sub_doc_path)
        composer.append(sub_doc)
    
    # Di chuyển tất cả thư mục "media" từ các file con vào gốc
    for folder in os.listdir(doc_folder):
        folder_path = os.path.join(doc_folder, folder)
        if os.path.isdir(folder_path) and folder.startswith('page_'):
            media_folder = os.path.join(folder_path, 'media')
            if os.path.exists(media_folder):
                for media_file in os.listdir(media_folder):
                    shutil.move(os.path.join(media_folder, media_file), os.path.join(doc_folder, media_file))
                shutil.rmtree(media_folder)
    
    # Lưu file hợp nhất
    composer.save(output_file)
    print(f"Hợp nhất hoàn tất: {output_file}")

# Xử lý PDF thành file Word
def process_pdf_to_word(pdf_path):
    # Xóa dữ liệu trong thư mục lưu trữ trước khi bắt đầu
    clear_folder(save_folder)
    
    # Lấy tên file gốc từ đường dẫn PDF
    input_filename = os.path.splitext(os.path.basename(pdf_path))[0]
    output_file = os.path.join(output_folder, f'{input_filename}_output.docx')

    # Chuyển đổi từng trang PDF thành hình ảnh
    images = convert_from_path(pdf_path)
    
    for i, image in enumerate(tqdm(images, desc="Pages processing")):
        # Chuyển đổi từ PIL Image sang OpenCV Image
        img = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
        process_image_to_word(img, i + 1)
    
    # Gộp các file Word
    merge_word_documents_with_images(save_folder, output_file)
    print(f"Hoàn thành! File Word hợp nhất lưu tại: {output_file}")

# Đường dẫn file PDF đầu vào
pdf_path = './inputs/pdf/example1.pdf'
process_pdf_to_word(pdf_path)

