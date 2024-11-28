# Copyright (c) 2020 PaddlePaddle Authors. All Rights Reserved.
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

# Update by KaiKenju (hiepdv.tb288@gmail.com)
import os
from copy import deepcopy
import cv2
import re
from docx import Document
from docx import shared
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches
from Recovery.table_process import HtmlToDocx
from tqdm import tqdm
from utils.logging import get_logger
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, MBartForConditionalGeneration, MBart50TokenizerFast
from transformers import pipeline
from Translation.pychecker import correct_title
from SpellCorrection.correct_spell import spell_checker
from SpellCorrection.correct_spell_vi import correct_spelling
from Model.model import models, AutoModelForSeq2SeqLM
logger = get_logger()

def load_model_and_tokenizer(lang):
    """
    Hàm tải model và tokenizer cho ngôn ngữ được yêu cầu nếu chưa được tải.
    :param lang: Mã ngôn ngữ ('vi', 'jp', ...)
    """
    if lang not in models:
        raise ValueError(f"Ngôn ngữ không được hỗ trợ: {lang}")

    config = models[lang]
    if config["model"] is None or config["tokenizer"] is None:  # Chỉ tải nếu chưa có
        print(f"Tải mô hình và tokenizer cho ngôn ngữ: {lang}")
        if config["use_lang_codes"]:
            config["tokenizer"] = MBart50TokenizerFast.from_pretrained(config["model_name"])
            config["model"] = MBartForConditionalGeneration.from_pretrained(config["model_name"])
        else:
            config["tokenizer"] = AutoTokenizer.from_pretrained(config["model_name"])
            config["model"] = AutoModelForSeq2SeqLM.from_pretrained(config["model_name"])

def translate(text, lang, device='cpu'):
    """
    Hàm dịch văn bản dựa trên ngôn ngữ và mô hình đã được cấu hình.
    :param text: Văn bản cần dịch.
    :param lang: Mã ngôn ngữ ('vi' hoặc 'jp').
    :return: Văn bản đã dịch.
    """
    if lang not in models:
        raise ValueError(f"Unsupported language: {lang}")
    
    # Lấy cấu hình mô hình
    load_model_and_tokenizer(lang)
    config = models[lang]
    tokenizer = config["tokenizer"]
    model = config["model"]

    # Xử lý cho mô hình có src_lang/target_lang
    if config["use_lang_codes"]:
        tokenizer.src_lang = config["src_lang"]  # Thiết lập ngôn ngữ nguồn
        inputs = tokenizer(text, return_tensors="pt", padding=True, truncation=True).to(device)
        outputs = model.generate(
            **inputs,
            forced_bos_token_id=tokenizer.lang_code_to_id[config["target_lang"]],
        )
    else:
        # Xử lý cho mô hình không dùng src_lang/target_lang (VietAI)
        inputs = tokenizer(text, return_tensors="pt", padding=True, truncation=True, max_length=512).to(device)
        outputs = model.generate(
            inputs.input_ids,
            max_length=512,
            num_beams=4,
            no_repeat_ngram_size=2,
            repetition_penalty=1.5,
            temperature=0.7,
        )
    
    # Giải mã kết quả dịch
    translated_text = tokenizer.batch_decode(outputs, skip_special_tokens=True)
    return translated_text[0].replace("vi: ", "").replace("vi", "")


def merge_lines_to_sentences(lines):
    merged_sentences = []
    current_sentence = ""

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # add current line with temp sentence
        current_sentence += " " + line

        if line.endswith((".", ";", "!", "?")):
            merged_sentences.append(current_sentence.strip())
            current_sentence = ""

    if current_sentence:
        merged_sentences.append(current_sentence.strip())

    return merged_sentences

def convert_info_docx(img, res, save_folder, img_name,lang, device='cpu'):
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.styles["Normal"].font.size = shared.Pt(12)

    flag = 1
    for i, region in enumerate(tqdm(res, desc="Converting regions to docx")):
        # if len(region["res"]) == 0:
        #     continue
        img_idx = region["img_idx"]
        if flag == 2 and region["layout"] == "single":
            section = doc.add_section(WD_SECTION.CONTINUOUS)
            section._sectPr.xpath("./w:cols")[0].set(qn("w:num"), "1")
            flag = 1
        elif flag == 1 and region["layout"] == "double":
            section = doc.add_section(WD_SECTION.CONTINUOUS)
            section._sectPr.xpath("./w:cols")[0].set(qn("w:num"), "2")
            flag = 2

        if str(region.get("type", "")).lower() == "figure":
            excel_save_folder = os.path.join(save_folder, img_name)
            if not os.path.exists(excel_save_folder):
                os.makedirs(excel_save_folder)

            # create input path
            bbox = region.get("bbox", "default_bbox")
            img_idx = region.get("img_idx", "default_idx")
            img_path = os.path.join(
                excel_save_folder, "{}_{}.jpg".format(bbox, img_idx)
            )

            # check input path
            if not os.path.exists(img_path):
                print("\n")
                print(f"The image doesn't exits: {img_path}")
                continue

            # check size image
            img = cv2.imread(img_path)
            if img is None or img.shape[0] < 10 or img.shape[1] < 10:
                print("\n")
                print(f"Image's too small or not exits: {img_path}")
                continue

            # add to word
            if os.path.exists(img_path):
                paragraph_pic = doc.add_paragraph()
                paragraph_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph_pic.add_run("")
                if flag == 1:
                    run.add_picture(img_path, width=Inches(5))
                elif flag == 2:
                    run.add_picture(img_path, width=Inches(2.6))
            else:
                print(f"Image does not exist at path: {img_path}")

        
        elif region["type"].lower() == "title":
            original_title = region["res"][0]["text"]
            corrected_title = correct_title(original_title)
            translated_title = translate(corrected_title, lang=lang, device=device)
            doc.add_heading(translated_title)
        elif region["type"].lower() == "table":
            parser = HtmlToDocx()
            parser.table_style = "TableGrid"
            parser.handle_table(region["res"]["html"], doc)
        elif region["type"] == "equation" or "latex" in region["res"]: # fix nốt với các công thức
            pass
        else:
            paragraph = doc.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            # # for i, line in enumerate(region["res"]):
            # #     if i == 0:
            # #         paragraph_format.first_line_indent = shared.Inches(0.25)
            # #     # original text : en
            # #     # text_run = paragraph.add_run(line["text"] + " ")
            # #     # text_run.font.size = shared.Pt(12)
                
            # #     # translate text
            # #     vietnamese_text = translate_to_vietnamese(line["text"])
            # #     text_run = paragraph.add_run(vietnamese_text + " ")
            # #     text_run.font.size = shared.Pt(12)
            #---------
            seen_texts = set()
            complete_paragraph = ""
            current_sentence = ""
            for i, line in enumerate(region["res"]):
                if i == 0:
                    paragraph_format.first_line_indent = shared.Inches(0.25)
                original_text = line["text"].strip()
                # print(f"OCR Output: {original_text}")
                corrected_text = spell_checker(original_text)[0]['generated_text']
                # print(f"Spell Correction: {corrected_text}")

                if corrected_text in seen_texts:
                    continue
                seen_texts.add(corrected_text)

                if current_sentence and not current_sentence.endswith(" "):
                    current_sentence += " "
                
                # Thêm văn bản hiện tại
                current_sentence +=  corrected_text
                if corrected_text.endswith((".", "!", "?")):
                    # corrected_text = spell_checker(original_text)[0]['generated_text']
                    
                    translated_text = translate(current_sentence.strip(), lang=lang,device=device)
                    if lang=="vi":
                         corrected_translated_text = correct_spelling(translated_text)
                    else:
                        print("KhÔNG Phải ngôn ngữ Vi nên ko check spell correction được")
                    print("\nOriginal:", current_sentence.strip())
                    print("\nTranslated&Correct:", corrected_translated_text)
                    print("------------------------------------")
                    complete_paragraph += current_sentence + " "
                    text_run = paragraph.add_run(corrected_translated_text + " ")
                    text_run.font.size = shared.Pt(12)
                    
                
                    current_sentence = ""

            if current_sentence.strip():
                # corrected_text = spell_checker(current_sentence)[0]['generated_text']
                
                translated_text = translate(current_sentence.strip(), lang=lang,device=device)
                if lang=="vi":
                    corrected_translated_text = correct_spelling(translated_text)
                else:
                    print("KhÔNG Phải ngôn ngữ Vi nên ko check spell correction được")
                # print("\nOriginal(Remaining):", current_sentence.strip())
                print("\nTranslated(Remaining):", corrected_translated_text)
                print("------------------------------------")
                complete_paragraph += current_sentence + " "
                text_run = paragraph.add_run(corrected_translated_text + " ")
                text_run.font.size = shared.Pt(12)

                
            # print all text in 1 phrase
            if complete_paragraph.strip():
                print("\nComplete Paragraph:")
                print(complete_paragraph)

    # save to docx
    docx_path = os.path.join(save_folder, "{}_ocr.docx".format(img_name))
    doc.save(docx_path)
    logger.info("docx save to {}".format(docx_path))


def sorted_layout_boxes(res, w):
    """
    Sort text boxes in order from top to bottom, left to right
    args:
        res(list):ppstructure results
    return:
        sorted results(list)
    """
    num_boxes = len(res)
    if num_boxes == 1:
        res[0]["layout"] = "single"
        return res

    sorted_boxes = sorted(res, key=lambda x: (x["bbox"][1], x["bbox"][0]))
    _boxes = list(sorted_boxes)

    new_res = []
    res_left = []
    res_right = []
    i = 0

    while True:
        if i >= num_boxes:
            break
        if i == num_boxes - 1:
            if (
                _boxes[i]["bbox"][1] > _boxes[i - 1]["bbox"][3]
                and _boxes[i]["bbox"][0] < w / 2
                and _boxes[i]["bbox"][2] > w / 2
            ):
                new_res += res_left
                new_res += res_right
                _boxes[i]["layout"] = "single"
                new_res.append(_boxes[i])
            else:
                if _boxes[i]["bbox"][2] > w / 2:
                    _boxes[i]["layout"] = "double"
                    res_right.append(_boxes[i])
                    new_res += res_left
                    new_res += res_right
                elif _boxes[i]["bbox"][0] < w / 2:
                    _boxes[i]["layout"] = "double"
                    res_left.append(_boxes[i])
                    new_res += res_left
                    new_res += res_right
            res_left = []
            res_right = []
            break
        elif _boxes[i]["bbox"][0] < w / 4 and _boxes[i]["bbox"][2] < 3 * w / 4:
            _boxes[i]["layout"] = "double"
            res_left.append(_boxes[i])
            i += 1
        elif _boxes[i]["bbox"][0] > w / 4 and _boxes[i]["bbox"][2] > w / 2:
            _boxes[i]["layout"] = "double"
            res_right.append(_boxes[i])
            i += 1
        else:
            new_res += res_left
            new_res += res_right
            _boxes[i]["layout"] = "single"
            new_res.append(_boxes[i])
            res_left = []
            res_right = []
            i += 1
    if res_left:
        new_res += res_left
    if res_right:
        new_res += res_right
    return new_res


